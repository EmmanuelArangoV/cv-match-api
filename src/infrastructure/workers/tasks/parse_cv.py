"""
Tarea Celery: descarga el CV de R2, extrae el contenido según el tipo de archivo
(PDF → imágenes vía PyMuPDF, DOCX → texto + imágenes embebidas, imágenes → directo),
llama a gpt-4o vision para extraer el perfil estructurado, genera el PDF normalizado
en estilo BBLABS y lo sube a R2.
"""
from __future__ import annotations

import base64
import io
import json
import uuid
import zipfile

import pymupdf as fitz  # PyMuPDF >= 1.24
from openai import OpenAI
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from src.config import settings
from src.infrastructure.ai.prompts import CV_EXTRACTION_PROMPT
from src.infrastructure.cv.pdf_renderer import render_normalized_cv
from src.infrastructure.storage.r2_client import download_file_sync, upload_file_sync
from src.infrastructure.workers.celery_app import celery_app

_engine = create_engine(settings.database_url_sync)
_SyncSession = sessionmaker(bind=_engine)

_INPUT_COST = 0.0000025
_OUTPUT_COST = 0.000010

# Extensiones reconocidas como imágenes directas
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".tiff", ".tif", ".bmp"}
_IMAGE_MIME = {
    ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".png": "image/png", ".webp": "image/webp",
    ".gif": "image/gif", ".tiff": "image/tiff", ".tif": "image/tiff",
    ".bmp": "image/bmp",
}


def _get_openai() -> OpenAI:
    return OpenAI(api_key=settings.openai_api_key)


# ─── Extractores por formato ───────────────────────────────────────────────────

def _pdf_to_content(pdf_bytes: bytes) -> list[dict]:
    """Convierte cada página del PDF a un bloque image_url para la API de visión."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    items: list[dict] = []
    for page in doc:
        mat = fitz.Matrix(2.0, 2.0)  # zoom 2× para mejor resolución
        pix = page.get_pixmap(matrix=mat)
        b64 = base64.b64encode(pix.tobytes("png")).decode()
        items.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{b64}", "detail": "high"},
        })
    doc.close()
    return items


def _docx_to_content(docx_bytes: bytes) -> list[dict]:
    """
    Extrae texto e imágenes embebidas de un DOCX.
    Devuelve bloques listos para la API de OpenAI:
      - un bloque 'text' con todo el texto plano del documento
      - un bloque 'image_url' por cada imagen embebida (hasta 10)
    """
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(docx_bytes))
    items: list[dict] = []

    # ── Texto plano ──────────────────────────────────────────────────────────
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                lines.append(row_text)

    if lines:
        items.append({"type": "text", "text": "DOCUMENTO (texto extraído):\n" + "\n".join(lines)})

    # ── Imágenes embebidas en el zip del DOCX ────────────────────────────────
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            media = [
                n for n in z.namelist()
                if n.startswith("word/media/") and not n.endswith("/")
            ]
            for name in media[:10]:  # limitamos a 10 imágenes por documento
                ext = "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""
                mime = _IMAGE_MIME.get(ext)
                if not mime:
                    continue
                b64 = base64.b64encode(z.read(name)).decode()
                items.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"},
                })
    except Exception:
        pass  # si falla la extracción de imágenes, continuamos solo con texto

    return items


def _image_to_content(image_bytes: bytes, ext: str) -> list[dict]:
    """Encoda una imagen directamente como bloque image_url."""
    mime = _IMAGE_MIME.get(ext, "image/jpeg")
    b64 = base64.b64encode(image_bytes).decode()
    return [{
        "type": "image_url",
        "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"},
    }]


def _prepare_content(file_bytes: bytes, r2_key: str) -> list[dict]:
    """Devuelve la lista de bloques de contenido según el tipo de archivo."""
    ext = ""
    if "." in r2_key:
        ext = "." + r2_key.rsplit(".", 1)[-1].lower()

    if ext == ".docx" or ext == ".doc":
        return _docx_to_content(file_bytes)
    if ext in _IMAGE_EXTS:
        return _image_to_content(file_bytes, ext)
    # Por defecto (PDF u otros): intentar como PDF
    return _pdf_to_content(file_bytes)


# ─── Llamada a OpenAI ──────────────────────────────────────────────────────────

def _call_openai(content_blocks: list[dict], client: OpenAI) -> tuple[dict, int, int]:
    """Llama a gpt-4o con el prompt de extracción + los bloques de contenido."""
    content: list[dict] = [{"type": "text", "text": CV_EXTRACTION_PROMPT}]
    content.extend(content_blocks)

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": content}],
        response_format={"type": "json_object"},
        max_tokens=4096,
    )

    result = json.loads(response.choices[0].message.content)
    return result, response.usage.prompt_tokens, response.usage.completion_tokens


# ─── Tarea Celery ──────────────────────────────────────────────────────────────

@celery_app.task(
    bind=True,
    max_retries=2,
    default_retry_delay=60,
    name="parse_cv",
)
def parse_cv(
    self,
    candidate_id: str,
    process_candidate_id: str,
    process_id: str,
) -> dict:
    from src.infrastructure.db.models import (
        Candidate,
        CandidateStatus,
        CostLog,
        OperationType,
        ProcessCandidate,
    )

    cand_uuid = uuid.UUID(candidate_id)
    pc_uuid = uuid.UUID(process_candidate_id)
    proc_uuid = uuid.UUID(process_id)

    with _SyncSession() as db:
        try:
            candidate: Candidate = db.get(Candidate, cand_uuid)
            pc: ProcessCandidate = db.get(ProcessCandidate, pc_uuid)

            if not candidate or not pc:
                return {"error": "Candidate or ProcessCandidate not found"}

            # Marcar en proceso
            pc.status = CandidateStatus.CV_PROCESSING.value
            db.commit()

            # Descargar el archivo de R2
            file_bytes = download_file_sync(candidate.cv_file_url)

            # Preparar bloques de contenido según el tipo de archivo
            content_blocks = _prepare_content(file_bytes, candidate.cv_file_url)

            if not content_blocks:
                raise ValueError(
                    f"No se pudo extraer contenido del archivo: {candidate.cv_file_url}"
                )

            # Llamar a OpenAI
            extracted, tokens_in, tokens_out = _call_openai(content_blocks, _get_openai())

            # Actualizar candidato con datos extraídos
            candidate.extracted_profile = extracted
            candidate.normalized_cv = extracted

            # Actualizar campos básicos si OpenAI los devolvió
            full_name: str = extracted.get("full_name", "")
            if full_name and " " in full_name:
                parts = full_name.split(" ", 1)
                candidate.name = parts[0][:100]
                candidate.last_name = parts[1][:100]
            elif full_name:
                candidate.name = full_name[:100]

            if extracted.get("email") and "@placeholder" in candidate.email:
                candidate.email = extracted["email"]

            if extracted.get("phone"):
                candidate.phone = extracted["phone"][:20]

            # Generar PDF normalizado en estilo BBLABS y subirlo a R2
            normalized_pdf_bytes = render_normalized_cv(extracted)
            original_key = candidate.cv_file_url
            # Genera siempre una key _normalized.pdf independientemente de la extensión original
            if "." in original_key.rsplit("/", 1)[-1]:
                normalized_key = original_key.rsplit(".", 1)[0] + "_normalized.pdf"
            else:
                normalized_key = original_key + "_normalized.pdf"
            upload_file_sync(normalized_key, normalized_pdf_bytes, "application/pdf")
            candidate.normalized_cv_url = normalized_key

            # Actualizar estado del proceso-candidato
            pc.status = CandidateStatus.MATCH_PENDING.value

            # Registrar costo
            estimated_cost = (tokens_in * _INPUT_COST) + (tokens_out * _OUTPUT_COST)
            cost_log = CostLog(
                process_id=proc_uuid,
                candidate_id=cand_uuid,
                operation_type=OperationType.CV_EXTRACTION.value,
                model_used="gpt-4o",
                tokens_input=tokens_in,
                tokens_output=tokens_out,
                estimated_cost=estimated_cost,
            )
            db.add(cost_log)
            db.commit()

            # Disparar automáticamente la tarea de match
            from src.infrastructure.workers.tasks.run_match import run_match
            run_match.delay(
                process_candidate_id=process_candidate_id,
                process_id=process_id,
            )

            # Disparar tarea de WhatsApp solo si las credenciales están configuradas
            if settings.meta_whatsapp_access_token and settings.meta_whatsapp_phone_number_id:
                from src.infrastructure.workers.tasks.whatsapp import send_whatsapp_consent
                send_whatsapp_consent.delay(process_candidate_id)

            return {
                "candidate_id": candidate_id,
                "status": "MATCH_PENDING",
                "tokens_in": tokens_in,
                "tokens_out": tokens_out,
                "estimated_cost_usd": round(estimated_cost, 6),
            }

        except Exception as exc:
            db.rollback()
            try:
                pc = db.get(ProcessCandidate, pc_uuid)
                if pc:
                    pc.status = CandidateStatus.CV_ERROR.value
                    db.commit()
            except Exception:
                pass
            raise self.retry(exc=exc)
